
import io
import re
import zipfile
import base64
from datetime import datetime
from typing import Dict, List, Optional

import numpy as np
import pandas as pd
import streamlit as st

# Optional exports
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

APP_TITLE = "منصة التحصيل الأكاديمي الذكي - مدرسة معيذر"
SCHOOL_NAME = "مدرسة معيذر الابتدائية للبنين"
ACADEMIC_YEAR = "2025-2026"

REQUIRED_COLUMNS = [
    "student_id", "student_name", "grade", "class_name", "subject", "teacher_name",
    "mid_first_15", "end_first_40", "mid_second_15", "end_year_40", "final_100"
]

COLUMN_ALIASES = {
    "student_id": ["student_id", "رقم الطالب", "الرقم", "id", "رقم"],
    "student_name": ["student_name", "اسم الطالب", "الاسم", "name", "الطالب"],
    "grade": ["grade", "الصف", "صف", "المستوى"],
    "class_name": ["class_name", "الشعبة", "الشعبة الصفية", "class", "الفصل"],
    "subject": ["subject", "المادة", "اسم المادة"],
    "teacher_name": ["teacher_name", "المعلم", "اسم المعلم", "teacher"],
    "mid_first_15": ["mid_first_15", "منتصف الفصل الاول", "منتصف أول", "منتصف الاول", "Mid First", "منتصف الفصل الأول"],
    "end_first_40": ["end_first_40", "نهاية الفصل الأول", "نهاية أول", "نهاية الاول", "End First"],
    "mid_second_15": ["mid_second_15", "منتصف الفصل الثاني", "منتصف ثاني", "Mid Second"],
    "end_year_40": ["end_year_40", "نهاية العام", "اختبار نهاية العام", "نهاية السنة", "End Year"],
    "final_100": ["final_100", "الدرجة الكلية 100", "النتيجة النهائية", "النهائية 100", "Final 100"],
}

TEACHER_COLUMNS = ["teacher_name", "subject", "grade", "class_name"]
CLASS_COLUMNS = ["grade", "class_name", "stage"]

st.set_page_config(page_title=APP_TITLE, layout="wide", page_icon="📊")

# ---------- Helpers ----------
def _clean_col(x):
    return str(x).strip().replace("\n", " ").replace("\r", " ")

def map_columns(df: pd.DataFrame) -> pd.DataFrame:
    original_cols = list(df.columns)
    cleaned_map = {_clean_col(c): c for c in original_cols}
    new = df.copy()
    rename = {}
    for std, aliases in COLUMN_ALIASES.items():
        for alias in aliases:
            alias_clean = _clean_col(alias)
            if alias_clean in cleaned_map:
                rename[cleaned_map[alias_clean]] = std
                break
    new = new.rename(columns=rename)
    for c in REQUIRED_COLUMNS:
        if c not in new.columns:
            new[c] = np.nan
    return new[REQUIRED_COLUMNS]

def to_numeric_safe(df, cols):
    for c in cols:
        df[c] = pd.to_numeric(df[c], errors="coerce")
    return df

def normalize_grade_label(x):
    s = str(x).strip()
    if s in ["5", "05", "خامس", "الخامس"]:
        return "خامس"
    if s in ["6", "06", "سادس", "السادس"]:
        return "سادس"
    if s.startswith("05") or s.startswith("5/"):
        return "خامس"
    if s.startswith("06") or s.startswith("6/"):
        return "سادس"
    return s

def prepare_grades(df: pd.DataFrame) -> pd.DataFrame:
    df = map_columns(df)
    df = to_numeric_safe(df, ["mid_first_15", "end_first_40", "mid_second_15", "end_year_40", "final_100"])
    df["grade"] = df["grade"].apply(normalize_grade_label)
    # If final_100 is blank, compute proportionally from the four exam windows.
    # Max = 15 + 40 + 15 + 40 = 110, then converted to 100.
    exam_sum = df[["mid_first_15", "end_first_40", "mid_second_15", "end_year_40"]].sum(axis=1, min_count=1)
    df["final_100"] = df["final_100"].fillna((exam_sum / 110) * 100)
    df["mid_first_pct"] = (df["mid_first_15"] / 15) * 100
    df["end_first_pct"] = (df["end_first_40"] / 40) * 100
    df["mid_second_pct"] = (df["mid_second_15"] / 15) * 100
    df["end_year_pct"] = (df["end_year_40"] / 40) * 100
    df["achievement_pct"] = df["final_100"].clip(0, 100)
    df["success"] = np.where(df["achievement_pct"] >= 50, "ناجح", "راسب")
    df["level"] = np.select(
        [df["achievement_pct"] >= 90, df["achievement_pct"] < 70],
        ["متميز", "متدني"],
        default="متوسط"
    )
    # Value added: all windows compared to Mid First as a fixed baseline.
    df["va_end_first_vs_mid_first"] = df["end_first_pct"] - df["mid_first_pct"]
    df["va_mid_second_vs_mid_first"] = df["mid_second_pct"] - df["mid_first_pct"]
    df["va_end_year_vs_mid_first"] = df["end_year_pct"] - df["mid_first_pct"]
    df["va_final_vs_mid_first"] = df["achievement_pct"] - df["mid_first_pct"]
    return df

def import_any_excel(files) -> pd.DataFrame:
    frames = []
    for f in files:
        xls = pd.ExcelFile(f)
        for sheet in xls.sheet_names:
            tmp = pd.read_excel(xls, sheet_name=sheet)
            if tmp.dropna(how="all").empty:
                continue
            frames.append(prepare_grades(tmp))
    if not frames:
        return pd.DataFrame(columns=REQUIRED_COLUMNS)
    return pd.concat(frames, ignore_index=True)

def metric_card(label, value, suffix=""):
    st.metric(label, f"{value}{suffix}")

def pct(x):
    if pd.isna(x):
        return 0
    return round(float(x), 2)

def safe_avg(series):
    return pct(pd.to_numeric(series, errors="coerce").mean())

def summary_by(df, group_cols):
    if df.empty:
        return pd.DataFrame()
    out = df.groupby(group_cols, dropna=False).agg(
        عدد_السجلات=("student_name", "count"),
        عدد_الطلاب=("student_id", pd.Series.nunique),
        متوسط_التحصيل=("achievement_pct", "mean"),
        نسبة_النجاح=("achievement_pct", lambda s: (s >= 50).mean() * 100),
        نسبة_التميز=("achievement_pct", lambda s: (s >= 90).mean() * 100),
        نسبة_التدني=("achievement_pct", lambda s: (s < 70).mean() * 100),
        القيمة_المضافة_نهاية_أول=("va_end_first_vs_mid_first", "mean"),
        القيمة_المضافة_منتصف_ثاني=("va_mid_second_vs_mid_first", "mean"),
        القيمة_المضافة_نهاية_العام=("va_end_year_vs_mid_first", "mean"),
        القيمة_المضافة_النهائية=("va_final_vs_mid_first", "mean"),
    ).reset_index()
    for c in out.columns:
        if c not in group_cols and c not in ["عدد_السجلات", "عدد_الطلاب"]:
            out[c] = out[c].round(2)
    return out

def recommendation(p):
    if p < 50:
        return "علاجي مكثف: خطة فردية عاجلة، حصص تقوية، اختبارات قصيرة أسبوعية، وتواصل مباشر مع ولي الأمر."
    if p < 70:
        return "علاجي: معالجة المهارات الأساسية، واجبات موجهة، متابعة أسبوعية، وإعادة قياس بعد 3 أسابيع."
    if p < 85:
        return "تحسيني: تدريبات إضافية، متابعة نصف شهرية، ورفع مستوى الأسئلة تدريجياً."
    if p < 90:
        return "تحفيزي: مهام تفكير عليا ومتابعة للوصول إلى التميز."
    return "إثرائي: مسابقات، مشاريع بحثية، نادي التميز، وأدوار قيادية للطالب."

def intervention_plan(p, va):
    if p < 50:
        priority, typ, dur = "عاجلة جداً", "علاجي مكثف", "4 أسابيع"
    elif p < 70:
        priority, typ, dur = "عاجلة", "علاجي", "3 أسابيع"
    elif p < 85:
        priority, typ, dur = "متوسطة", "تحسيني", "أسبوعان"
    elif p < 90:
        priority, typ, dur = "منخفضة", "تحفيزي", "أسبوعان"
    else:
        priority, typ, dur = "إثراء", "إثرائي", "مستمر"
    return {
        "نوع التدخل": typ,
        "الأولوية": priority,
        "الإجراءات": recommendation(p),
        "المسؤول": "معلم المادة + لجنة التحصيل الأكاديمي",
        "المدة": dur,
        "مؤشر الأثر": "تحسن القيمة المضافة مقارنة بمنتصف الفصل الأول وارتفاع النتيجة النهائية."
    }

def build_interventions(df):
    rows = []
    for _, r in df.iterrows():
        plan = intervention_plan(r["achievement_pct"], r["va_final_vs_mid_first"])
        rows.append({
            "الطالب": r["student_name"],
            "الصف": r["grade"],
            "الشعبة": r["class_name"],
            "المادة": r["subject"],
            "المعلم": r["teacher_name"],
            "النسبة النهائية": round(r["achievement_pct"], 2),
            "القيمة المضافة النهائية": round(r["va_final_vs_mid_first"], 2),
            **plan
        })
    return pd.DataFrame(rows)

def to_excel_bytes(sheets: Dict[str, pd.DataFrame]) -> bytes:
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
        for name, data in sheets.items():
            data.to_excel(writer, index=False, sheet_name=name[:31])
            ws = writer.sheets[name[:31]]
            for i, col in enumerate(data.columns):
                ws.set_column(i, i, min(max(len(str(col)) + 4, 14), 35))
    return bio.getvalue()

def docx_report(title, tables: Dict[str, pd.DataFrame], leadership: Dict[str, str]) -> bytes:
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"{SCHOOL_NAME} - العام الأكاديمي {ACADEMIC_YEAR}")
    if leadership:
        doc.add_paragraph(f"مدير المدرسة: {leadership.get('principal','')}")
        doc.add_paragraph(f"نائب المدير للشؤون الأكاديمية: {leadership.get('academic_vp','')}")
    for name, df in tables.items():
        doc.add_heading(name, level=1)
        if df.empty:
            doc.add_paragraph("لا توجد بيانات.")
            continue
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        for i, col in enumerate(df.columns):
            table.rows[0].cells[i].text = str(col)
        for _, row in df.head(200).iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def pdf_report(title, tables: Dict[str, pd.DataFrame], leadership: Dict[str, str]) -> bytes:
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=landscape(A4), rightMargin=20, leftMargin=20, topMargin=20, bottomMargin=20)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Paragraph(f"{SCHOOL_NAME} - {ACADEMIC_YEAR}", styles["Normal"])]
    if leadership:
        story.append(Paragraph(f"Principal: {leadership.get('principal','')} | Academic VP: {leadership.get('academic_vp','')}", styles["Normal"]))
    story.append(Spacer(1, 12))
    for name, df in tables.items():
        story.append(Paragraph(name, styles["Heading2"]))
        if df.empty:
            story.append(Paragraph("No data", styles["Normal"]))
            continue
        small = df.head(40).astype(str)
        data = [list(small.columns)] + small.values.tolist()
        tbl = Table(data, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
            ("FONTSIZE", (0,0), (-1,-1), 7),
            ("ALIGN", (0,0), (-1,-1), "CENTER"),
        ]))
        story.append(tbl)
        story.append(PageBreak())
    doc.build(story)
    return bio.getvalue()

def print_button():
    st.markdown("""
    <script>
    function printPage(){window.print();}
    </script>
    <button onclick="window.print()" style="padding:10px 18px;border-radius:8px;border:1px solid #999;">
        🖨️ طباعة الصفحة
    </button>
    """, unsafe_allow_html=True)

# ---------- UI ----------
st.title("📊 منصة التحصيل الأكاديمي الذكي")
st.caption("مدرسة معيذر الابتدائية للبنين | الصف الخامس والسادس | 2025-2026")

with st.sidebar:
    st.header("⚙️ إعدادات")
    principal = st.text_input("مدير المدرسة")
    academic_vp = st.text_input("نائب المدير للشؤون الأكاديمية", value="حيدر يحيى الشرمان")
    prepared_by = st.text_input("إعداد", value="حيدر يحيى الشرمان")
    st.divider()
    st.subheader("📥 الاستيراد")
    uploaded = st.file_uploader("استيراد الدرجات حسب المواعيد", type=["xlsx", "xls"], accept_multiple_files=True)
    teachers_file = st.file_uploader("استيراد بيانات المعلمين", type=["xlsx", "xls"], key="teachers")
    classes_file = st.file_uploader("استيراد بيانات الصفوف والشعب", type=["xlsx", "xls"], key="classes")
    st.caption("يمكن أيضاً إدخال بيانات معلم/صف يدوياً من تبويب الإدخال اليدوي.")

leadership = {"principal": principal, "academic_vp": academic_vp, "prepared_by": prepared_by}

if uploaded:
    df = import_any_excel(uploaded)
else:
    df = pd.DataFrame(columns=REQUIRED_COLUMNS)

# Teacher/class imports
teachers_df = pd.DataFrame(columns=TEACHER_COLUMNS)
classes_df = pd.DataFrame(columns=CLASS_COLUMNS)
if teachers_file:
    raw = pd.read_excel(teachers_file)
    teachers_df = raw.rename(columns={
        "اسم المعلم": "teacher_name", "المعلم": "teacher_name",
        "المادة": "subject", "الصف": "grade", "الشعبة": "class_name"
    })
if classes_file:
    raw = pd.read_excel(classes_file)
    classes_df = raw.rename(columns={
        "الصف": "grade", "الشعبة": "class_name", "المرحلة": "stage"
    })

tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
    "🏠 Dashboard", "📥 إدخال/استيراد", "📊 تحليلات", "👨‍🎓 تقارير الطلاب",
    "⚠️ التدخل الذكي", "🎯 أثر البرامج", "📤 تصدير وطباعة"
])

with tab1:
    if df.empty:
        st.info("ارفع ملف الدرجات أو استخدم قالب البيانات من تبويب إدخال/استيراد.")
    else:
        c1, c2, c3, c4, c5 = st.columns(5)
        c1.metric("عدد الطلاب", df["student_id"].nunique())
        c2.metric("متوسط التحصيل", f"{safe_avg(df['achievement_pct'])}%")
        c3.metric("نسبة النجاح", f"{pct((df['achievement_pct'] >= 50).mean()*100)}%")
        c4.metric("نسبة التميز", f"{pct((df['achievement_pct'] >= 90).mean()*100)}%")
        c5.metric("نسبة التدني", f"{pct((df['achievement_pct'] < 70).mean()*100)}%")
        st.subheader("القيمة المضافة مقارنة بمنتصف الفصل الأول")
        va_cols = ["va_end_first_vs_mid_first", "va_mid_second_vs_mid_first", "va_end_year_vs_mid_first", "va_final_vs_mid_first"]
        st.bar_chart(df[va_cols].mean().rename({
            "va_end_first_vs_mid_first":"نهاية أول",
            "va_mid_second_vs_mid_first":"منتصف ثاني",
            "va_end_year_vs_mid_first":"نهاية العام",
            "va_final_vs_mid_first":"النهائية"
        }))
        st.subheader("توزيع المستويات")
        st.bar_chart(df["level"].value_counts())

with tab2:
    st.subheader("قالب الاستيراد المعتمد")
    template_df = pd.DataFrame(columns=[
        "student_id","student_name","grade","class_name","subject","teacher_name",
        "mid_first_15","end_first_40","mid_second_15","end_year_40","final_100"
    ])
    st.download_button("تحميل قالب الدرجات Excel", to_excel_bytes({"grades_template": template_df}), "grades_template.xlsx")
    st.download_button("تحميل قالب المعلمين Excel", to_excel_bytes({"teachers": pd.DataFrame(columns=TEACHER_COLUMNS)}), "teachers_template.xlsx")
    st.download_button("تحميل قالب الصفوف Excel", to_excel_bytes({"classes": pd.DataFrame(columns=CLASS_COLUMNS)}), "classes_template.xlsx")
    st.divider()
    st.subheader("إدخال يدوي سريع")
    with st.form("manual"):
        cols = st.columns(5)
        sid = cols[0].text_input("رقم الطالب")
        name = cols[1].text_input("اسم الطالب")
        grade = cols[2].selectbox("الصف", ["خامس", "سادس"])
        cls = cols[3].text_input("الشعبة")
        subj = cols[4].text_input("المادة")
        cols2 = st.columns(5)
        teacher = cols2[0].text_input("المعلم")
        mf = cols2[1].number_input("منتصف أول /15", min_value=0.0, max_value=15.0, value=0.0)
        ef = cols2[2].number_input("نهاية أول /40", min_value=0.0, max_value=40.0, value=0.0)
        ms = cols2[3].number_input("منتصف ثاني /15", min_value=0.0, max_value=15.0, value=0.0)
        ey = cols2[4].number_input("نهاية العام /40", min_value=0.0, max_value=40.0, value=0.0)
        final100 = st.number_input("الدرجة الكلية /100 (اختياري، إذا تركتها صفر يحسبها النظام)", min_value=0.0, max_value=100.0, value=0.0)
        submitted = st.form_submit_button("إضافة السجل للمعاينة")
        if submitted:
            manual_df = prepare_grades(pd.DataFrame([{
                "student_id": sid, "student_name": name, "grade": grade, "class_name": cls, "subject": subj,
                "teacher_name": teacher, "mid_first_15": mf, "end_first_40": ef, "mid_second_15": ms,
                "end_year_40": ey, "final_100": np.nan if final100 == 0 else final100
            }]))
            st.dataframe(manual_df, use_container_width=True)

with tab3:
    if not df.empty:
        st.subheader("تحليل حسب الصف")
        st.dataframe(summary_by(df, ["grade"]), use_container_width=True)
        st.subheader("تحليل حسب الشعبة")
        st.dataframe(summary_by(df, ["grade", "class_name"]), use_container_width=True)
        st.subheader("تحليل حسب المادة")
        st.dataframe(summary_by(df, ["subject"]), use_container_width=True)
        st.subheader("تحليل حسب المعلم")
        st.dataframe(summary_by(df, ["teacher_name"]), use_container_width=True)
        st.subheader("تحليل تفصيلي")
        st.dataframe(df, use_container_width=True)
    else:
        st.warning("لا توجد بيانات.")

with tab4:
    if not df.empty:
        student = st.selectbox("اختر الطالب", sorted(df["student_name"].dropna().unique()))
        sdf = df[df["student_name"] == student].copy()
        sdf["التوصية"] = sdf["achievement_pct"].apply(recommendation)
        display_cols = ["subject","teacher_name","mid_first_15","end_first_40","mid_second_15","end_year_40","final_100","achievement_pct","level","va_final_vs_mid_first","التوصية"]
        st.dataframe(sdf[display_cols], use_container_width=True)
        st.download_button("تصدير تقرير الطالب Word", docx_report(f"تقرير الطالب: {student}", {"تقرير الطالب": sdf[display_cols]}, leadership), f"{student}_report.docx")
        st.download_button("تصدير تقرير الطالب PDF", pdf_report(f"تقرير الطالب: {student}", {"تقرير الطالب": sdf[display_cols]}, leadership), f"{student}_report.pdf")
    else:
        st.warning("لا توجد بيانات.")

with tab5:
    if not df.empty:
        interventions = build_interventions(df)
        st.subheader("لوحة التدخل الذكي")
        c1, c2, c3 = st.columns(3)
        c1.metric("خطط علاجية", int(interventions["نوع التدخل"].str.contains("علاجي").sum()))
        c2.metric("خطط إثرائية", int(interventions["نوع التدخل"].str.contains("إثرائي").sum()))
        c3.metric("حالات عاجلة", int(interventions["الأولوية"].str.contains("عاجلة").sum()))
        st.dataframe(interventions, use_container_width=True)
    else:
        st.warning("لا توجد بيانات.")

with tab6:
    st.subheader("قياس أثر برنامج أكاديمي")
    if not df.empty:
        program_name = st.text_input("اسم البرنامج", "برنامج التقوية / الإثراء")
        subject_filter = st.selectbox("المادة", ["الكل"] + sorted(df["subject"].dropna().unique().tolist()))
        grade_filter = st.selectbox("الصف", ["الكل"] + sorted(df["grade"].dropna().unique().tolist()))
        students = sorted(df["student_name"].dropna().unique())
        sample_students = st.multiselect("طلاب عينة البرنامج", students)
        calc = st.button("حساب أثر البرنامج")
        if calc:
            work = df.copy()
            if subject_filter != "الكل": work = work[work["subject"] == subject_filter]
            if grade_filter != "الكل": work = work[work["grade"] == grade_filter]
            sample = work[work["student_name"].isin(sample_students)]
            comp = work[~work["student_name"].isin(sample_students)]
            sample_avg = safe_avg(sample["achievement_pct"])
            comp_avg = safe_avg(comp["achievement_pct"])
            diff = round(sample_avg - comp_avg, 2)
            va = safe_avg(sample["va_final_vs_mid_first"])
            if diff >= 10:
                status = "ناجح جداً - يوصى بالتعميم"
            elif diff >= 5:
                status = "ناجح - يوصى بالاستمرار"
            elif diff >= 2:
                status = "مقبول - يحتاج تطوير"
            elif diff > -2:
                status = "أثر غير واضح"
            else:
                status = "يحتاج مراجعة"
            impact_df = pd.DataFrame([{
                "اسم البرنامج": program_name,
                "عدد العينة": sample["student_id"].nunique(),
                "عدد المقارنة": comp["student_id"].nunique(),
                "متوسط العينة": sample_avg,
                "متوسط بقية الطلاب": comp_avg,
                "الفارق": diff,
                "القيمة المضافة للعينة": va,
                "الحكم": status
            }])
            st.dataframe(impact_df, use_container_width=True)
    else:
        st.warning("لا توجد بيانات.")

with tab7:
    st.subheader("تصدير وطباعة")
    print_button()
    if not df.empty:
        sheets = {
            "البيانات": df,
            "حسب الصف": summary_by(df, ["grade"]),
            "حسب الشعبة": summary_by(df, ["grade", "class_name"]),
            "حسب المادة": summary_by(df, ["subject"]),
            "حسب المعلم": summary_by(df, ["teacher_name"]),
            "التدخلات": build_interventions(df),
        }
        st.download_button("📊 تصدير جميع الجداول Excel", to_excel_bytes(sheets), "muaither_full_analysis.xlsx")
        st.download_button("📄 تصدير التقرير الشامل Word", docx_report("التقرير الشامل للتحصيل الأكاديمي", sheets, leadership), "muaither_report.docx")
        st.download_button("📑 تصدير التقرير الشامل PDF", pdf_report("التقرير الشامل للتحصيل الأكاديمي", sheets, leadership), "muaither_report.pdf")
        # ZIP of all student pdfs
        zip_bio = io.BytesIO()
        with zipfile.ZipFile(zip_bio, "w", zipfile.ZIP_DEFLATED) as zf:
            for student in sorted(df["student_name"].dropna().unique()):
                sdf = df[df["student_name"] == student].copy()
                sdf["التوصية"] = sdf["achievement_pct"].apply(recommendation)
                content = pdf_report(f"تقرير الطالب: {student}", {"تقرير الطالب": sdf}, leadership)
                zf.writestr(f"{student}_report.pdf", content)
        st.download_button("🗂️ تصدير جميع تقارير الطلاب PDF ZIP", zip_bio.getvalue(), "all_student_reports_pdf.zip")
    else:
        st.warning("لا توجد بيانات للتصدير.")
